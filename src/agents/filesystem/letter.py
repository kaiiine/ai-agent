"""
Génère une lettre de motivation en DOCX + PDF à partir du texte brut de la lettre.
"""
from __future__ import annotations

import re
import subprocess
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ACCENT_COLOR = RGBColor(0xE8, 0x7C, 0x1E)   # orange proche du thème
TEXT_COLOR   = RGBColor(0x1A, 0x1A, 0x2E)   # presque noir
GREY_COLOR   = RGBColor(0x6B, 0x6B, 0x6B)   # gris medium
FONT_NAME    = "Calibri"
OUTPUT_DIR   = Path.home() / "Documents" / "CV"


def _add_rule(doc: Document, color: RGBColor = ACCENT_COLOR, thickness_pt: int = 2) -> None:
    """Ajoute une ligne horizontale colorée."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness_pt * 4))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), str(color))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_run(run, text: str, size: int, bold: bool = False,
             color: RGBColor = TEXT_COLOR, italic: bool = False) -> None:
    run.text = text
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _para(doc: Document, text: str, size: int = 11, bold: bool = False,
          color: RGBColor = TEXT_COLOR, align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before: int = 0, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.alignment   = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run()
    _set_run(run, text, size, bold, color)


def _extract_header(text: str) -> tuple[str, str, str]:
    """Extrait objet, salutation, et le corps de la lettre."""
    lines = text.strip().splitlines()
    objet = ""
    salut = ""
    body_start = 0

    for i, line in enumerate(lines):
        if line.strip().lower().startswith("objet"):
            objet = line.strip()
            continue
        if line.strip().lower().startswith(("madame", "monsieur", "madame,")):
            salut = line.strip()
            body_start = i + 1
            break

    body = "\n".join(lines[body_start:]).strip()
    return objet, salut, body


def _split_paragraphs(body: str) -> list[str]:
    """Découpe le corps en paragraphes (ligne vide = séparateur)."""
    paras = [p.strip() for p in re.split(r"\n{2,}", body)]
    return [p for p in paras if p]


def generate_docx(letter_text: str, candidate_name: str = "Quentin Dufour",
                  company: str = "") -> Path:
    """
    Génère un fichier .docx professionnel à partir du texte brut de la lettre.
    Retourne le chemin du fichier créé.
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    slug = re.sub(r"[^a-z0-9]+", "-", company.lower()).strip("-") if company else "lettre"
    date_str = datetime.today().strftime("%Y-%m-%d")
    filename = OUTPUT_DIR / f"lettre_{slug}_{date_str}.docx"

    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.5)

    # ── EN-TÊTE : Nom + date ──────────────────────────────────────────
    header_p = doc.add_paragraph()
    header_p.paragraph_format.space_before = Pt(0)
    header_p.paragraph_format.space_after  = Pt(2)

    name_run = header_p.add_run(candidate_name + "    ")
    _set_run(name_run, candidate_name + "    ", size=15, bold=True, color=ACCENT_COLOR)

    date_run = header_p.add_run(f"le {datetime.today().strftime('%d %B %Y')}")
    _set_run(date_run, f"le {datetime.today().strftime('%d %B %Y')}", size=10, color=GREY_COLOR, italic=True)

    _add_rule(doc)
    _para(doc, "", size=4, space_after=8)   # espace après la ligne

    # ── OBJET ────────────────────────────────────────────────────────
    objet, salut, body = _extract_header(letter_text)

    if objet:
        _para(doc, objet, size=10, bold=True, color=TEXT_COLOR, space_after=14)

    # ── SALUTATION ───────────────────────────────────────────────────
    if salut:
        _para(doc, salut, size=11, space_before=4, space_after=14)

    # ── CORPS ────────────────────────────────────────────────────────
    paragraphs = _split_paragraphs(body)
    for i, para in enumerate(paragraphs):
        # Ignorer "Cordialement" et la signature — traités séparément
        if para.lower().startswith("cordialement"):
            break
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(10)
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(para)
        _set_run(run, para, size=11, color=TEXT_COLOR)

    # ── CLÔTURE ──────────────────────────────────────────────────────
    _para(doc, "", space_after=14)
    _para(doc, "Cordialement,", size=11, space_after=28)

    # Ligne de signature
    _add_rule(doc, color=ACCENT_COLOR, thickness_pt=1)
    _para(doc, candidate_name, size=11, bold=True, color=ACCENT_COLOR, space_before=6)

    doc.save(str(filename))
    return filename


def docx_to_pdf(docx_path: Path) -> Path:
    """Convertit le .docx en PDF via LibreOffice headless."""
    pdf_path = docx_path.with_suffix(".pdf")
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(docx_path.parent), str(docx_path)],
        check=True,
        capture_output=True,
        timeout=30,
    )
    return pdf_path
