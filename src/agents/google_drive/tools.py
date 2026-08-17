"""Google Drive — chercher, lire et gérer les fichiers.

Le module le mieux construit de la revue Workspace, et il portait quand même
quatre défauts, tous corrigés ici :

  1. une APOSTROPHE cassait la recherche. La requête interpolait le nom dans une
     chaîne entre quotes simples en n'échappant que les guillemets doubles, donc
     « Compte-rendu d'équipe » produisait `name contains 'Compte-rendu d'équipe'` :
     trois apostrophes, requête malformée, HTTP 400. En français, un nom de
     fichier sur trois est concerné ;
  2. les DRIVE PARTAGÉS étaient invisibles : aucun appel ne passait
     `includeItemsFromAllDrives`, donc un Drive d'équipe n'apparaissait jamais ;
  3. les PDF et les Google Slides n'étaient pas lisibles, alors que le PDF est le
     format le plus courant d'un Drive ;
  4. `permanently=True` supprimait définitivement, avec pour seule garde une
     phrase dans la docstring — « demander confirmation avant ». Une docstring
     n'est pas un garde-fou.
"""
from __future__ import annotations

import io
from typing import Any, Dict, Optional

from googleapiclient.errors import HttpError
from langchain_core.tools import tool

from src.infra.google_auth import get_drive_service

#: Passé à chaque appel : sans lui, un fichier d'un Drive d'équipe est traité
#: comme inexistant, ce qui ressemble à un problème de droits alors que c'est un
#: paramètre manquant.
_PARTAGES = {"supportsAllDrives": True}
_PARTAGES_LISTE = {"includeItemsFromAllDrives": True, "supportsAllDrives": True}

_DOC = "application/vnd.google-apps.document"
_FEUILLE = "application/vnd.google-apps.spreadsheet"
_PRESENTATION = "application/vnd.google-apps.presentation"
_PDF = "application/pdf"
_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _file_url(file_id: str) -> str:
    return f"https://drive.google.com/file/d/{file_id}/view"


def _echapper(valeur: str) -> str:
    """Échappe une valeur pour une requête Drive entre quotes SIMPLES.

    La syntaxe de recherche Drive attend des chaînes entre quotes simples, et
    veut un antislash devant l'antislash et devant la quote. L'ancien code
    n'échappait que le guillemet double — inutile ici, et il laissait passer
    l'apostrophe, qui est précisément le caractère qui casse la requête.
    """
    return valeur.replace("\\", "\\\\").replace("'", "\\'")


def _telecharger(svc, file_id: str) -> bytes:
    from googleapiclient.http import MediaIoBaseDownload

    tampon = io.BytesIO()
    telechargeur = MediaIoBaseDownload(
        tampon, svc.files().get_media(fileId=file_id, **_PARTAGES))
    fini = False
    while not fini:
        _, fini = telechargeur.next_chunk()
    return tampon.getvalue()


def _exporter(svc, file_id: str, mime: str) -> str:
    brut = svc.files().export(fileId=file_id, mimeType=mime).execute()
    return brut.decode("utf-8", errors="replace") if isinstance(brut, bytes) else brut


def _texte_de_pdf(octets: bytes) -> str:
    """Réutilise l'extracteur déjà réglé du projet, plutôt qu'un second.

    `src/ui/attachments._extract_pdf` gère déjà les caractères espacés et pose
    des marqueurs de page. En écrire un deuxième ici garantirait qu'ils divergent.
    """
    import tempfile
    from pathlib import Path

    from src.ui.attachments import _extract_pdf

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(octets)
        chemin = Path(f.name)
    try:
        return _extract_pdf(chemin)
    finally:
        chemin.unlink(missing_ok=True)


def _texte_de_docx(octets: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(octets))
    morceaux = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for rang in table.rows:
            cellules = [c.text.strip() for c in rang.cells]
            if any(cellules):
                morceaux.append("| " + " | ".join(cellules) + " |")
    return "\n".join(morceaux)


@tool("drive_find_file_id")
def drive_find_file_id(name: str, exact: bool = False) -> Dict[str, Any]:
    """
    Recherche des fichiers Google Drive par nom et retourne leurs identifiants.

    Utilise ce tool quand l'utilisateur veut :
    - trouver un fichier sur Google Drive par son nom
    - localiser un document, une feuille ou une présentation Drive
    - obtenir l'ID d'un fichier Drive pour d'autres opérations

    Mots-clés : Google Drive, fichier, trouver, chercher document, ID fichier, cloud, Drive

    Les Drive partagés sont inclus. Si plusieurs résultats, les lister et demander
    à l'utilisateur lequel il veut.

    Args:
        name: nom exact ou fragment à chercher (les apostrophes sont gérées)
        exact: True = nom exactement identique, False = contient le fragment
    Returns:
        {"status": "ok", "matches": [{"id", "name", "mimeType", "modifiedTime", "url"}, ...]}
        {"status": "empty"} si aucun résultat
    """
    svc = get_drive_service()
    try:
        motif = _echapper(name)
        q = (f"name = '{motif}' and trashed = false" if exact
             else f"name contains '{motif}' and trashed = false")
        resp = svc.files().list(
            q=q, spaces="drive",
            fields="files(id,name,mimeType,modifiedTime,driveId)",
            pageSize=50, **_PARTAGES_LISTE,
        ).execute()
        files = resp.get("files", [])
        if not files:
            return {"status": "empty", "matches": []}
        matches = [
            {
                "id": f["id"],
                "name": f["name"],
                "mimeType": f.get("mimeType"),
                "modifiedTime": f.get("modifiedTime"),
                "drive_partage": bool(f.get("driveId")),
                "url": _file_url(f["id"]),
            }
            for f in files
        ]
        return {"status": "ok", "count": len(matches), "matches": matches}
    except HttpError as e:
        return {"status": "error", "error": str(e)}


@tool("drive_read_file")
def drive_read_file(file_id: str) -> Dict[str, Any]:
    """
    Lit le contenu d'un fichier Drive : Docs, Sheets, Slides, PDF, Word, texte.

    Utilise ce tool quand l'utilisateur veut :
    - lire le contenu d'un Google Doc, d'un PDF ou d'un fichier Drive
    - accéder au texte d'un document partagé
    - voir ce qu'il y a dans une feuille Google Sheets ou une présentation

    Mots-clés : lire Drive, Google Docs, PDF, Slides, Sheets, contenu fichier, document cloud

    Args:
        file_id: ID du fichier Drive
    Returns:
        {"status": "ok", "name": "...", "content": "...", "mime_type": "..."}
        {"status": "unsupported", ...} si le type ne contient pas de texte
    """
    svc = get_drive_service()
    try:
        meta = svc.files().get(
            fileId=file_id, fields="id,name,mimeType,size", **_PARTAGES).execute()
        mime = meta.get("mimeType", "")
        name = meta.get("name", "")

        if mime == _DOC:
            texte, limite = _exporter(svc, file_id, "text/plain"), 50_000
        elif mime == _FEUILLE:
            texte, limite = _exporter(svc, file_id, "text/csv"), 20_000
        elif mime == _PRESENTATION:
            # Une présentation s'exporte en texte : les titres et les puces de
            # chaque diapositive, dans l'ordre. C'était rendu « non supporté ».
            texte, limite = _exporter(svc, file_id, "text/plain"), 50_000
        elif mime == _PDF:
            texte, limite = _texte_de_pdf(_telecharger(svc, file_id)), 50_000
        elif mime == _DOCX:
            texte, limite = _texte_de_docx(_telecharger(svc, file_id)), 50_000
        elif mime.startswith("text/") or mime in ("application/json", "text/csv"):
            texte = _telecharger(svc, file_id).decode("utf-8", errors="replace")
            limite = 50_000
        else:
            return {
                "status": "unsupported",
                "name": name,
                "mime_type": mime,
                "error": f"Pas de texte à extraire d'un {mime}. "
                         "Types lus : Docs, Sheets, Slides, PDF, Word, texte.",
            }

        return {
            "status": "ok",
            "name": name,
            "mime_type": mime,
            "content": texte[:limite],
            "tronque": len(texte) > limite,
        }
    except HttpError as e:
        if e.resp is not None and e.resp.status == 404:
            return {"status": "not_found", "file_id": file_id,
                    "error": "Fichier introuvable ou accès refusé"}
        return {"status": "error", "error": str(e)}
    except Exception as e:                                       # noqa: BLE001
        # Un PDF illisible ou un docx corrompu ne doit pas remonter en exception :
        # l'agent doit pouvoir le dire et passer à autre chose.
        return {"status": "error", "file_id": file_id,
                "error": f"Extraction impossible : {e}"}


@tool("drive_list_files")
def drive_list_files(q: Optional[str] = None, page_size: int = 50,
                     page_token: Optional[str] = None) -> Dict[str, Any]:
    """
    Liste les fichiers disponibles sur Google Drive avec filtre optionnel.

    Utilise ce tool quand l'utilisateur veut :
    - parcourir son Google Drive
    - lister tous ses fichiers cloud
    - filtrer par type de fichier dans Drive

    Mots-clés : Google Drive, liste fichiers, parcourir, cloud, documents en ligne

    Les Drive partagés sont inclus.

    Args:
        q: requête Drive optionnelle (ex: "mimeType='application/pdf'")
        page_size: nombre max de résultats (1 à 1000)
        page_token: pour pagination
    """
    svc = get_drive_service()
    try:
        resp = svc.files().list(
            q=q or "trashed = false", spaces="drive",
            fields="nextPageToken,files(id,name,mimeType,modifiedTime,driveId)",
            pageSize=max(1, min(page_size, 1000)),
            pageToken=page_token, **_PARTAGES_LISTE,
        ).execute()
        files = [
            {
                "id": f["id"],
                "name": f["name"],
                "mimeType": f.get("mimeType"),
                "modifiedTime": f.get("modifiedTime"),
                "drive_partage": bool(f.get("driveId")),
                "url": _file_url(f["id"]),
            }
            for f in resp.get("files", [])
        ]
        return {"status": "ok", "count": len(files), "files": files,
                "next_page_token": resp.get("nextPageToken")}
    except HttpError as e:
        return {"status": "error", "error": str(e)}


@tool("drive_delete_file")
def drive_delete_file(file_id: str, permanently: bool = False,
                      confirmer_nom: str = "") -> Dict[str, Any]:
    """
    Met un fichier Drive à la corbeille, ou le supprime définitivement.

    Utilise ce tool quand l'utilisateur veut :
    - supprimer un fichier de son Google Drive
    - mettre un document à la corbeille Drive

    Mots-clés : supprimer Drive, effacer fichier, corbeille, Google Drive, delete

    La corbeille est réversible et suffit presque toujours. La suppression
    DÉFINITIVE exige en plus `confirmer_nom` égal au nom exact du fichier : c'est
    une garde réelle, là où la consigne « demander confirmation » n'était qu'une
    phrase dans cette docstring. Récupère le nom avec drive_get_file_metadata et
    fais-le valider par l'utilisateur avant de le recopier ici.

    Args:
        file_id: ID du fichier
        permanently: True = suppression définitive (IRRÉVERSIBLE), False = corbeille
        permanently exige confirmer_nom : nom exact du fichier, recopié
    """
    svc = get_drive_service()
    try:
        if permanently:
            meta = svc.files().get(fileId=file_id, fields="name", **_PARTAGES).execute()
            reel = meta.get("name", "")
            if confirmer_nom.strip() != reel:
                return {
                    "status": "confirmation_requise",
                    "file_id": file_id,
                    "nom_reel": reel,
                    "error": "Suppression définitive refusée : `confirmer_nom` doit "
                             f"valoir exactement « {reel} ». Fais valider ce nom par "
                             "l'utilisateur, puis rappelle ce tool avec.",
                }
            svc.files().delete(fileId=file_id, **_PARTAGES).execute()
            return {"status": "ok", "supprime": reel,
                    "message": f"« {reel} » supprimé DÉFINITIVEMENT."}

        svc.files().update(fileId=file_id, body={"trashed": True}, **_PARTAGES).execute()
        return {"status": "ok", "file_id": file_id,
                "message": "Fichier mis à la corbeille (réversible)."}
    except HttpError as e:
        if e.resp is not None and e.resp.status == 404:
            return {"status": "not_found", "file_id": file_id,
                    "error": "Fichier introuvable ou accès refusé"}
        return {"status": "error", "error": str(e)}


@tool("drive_get_file_metadata")
def drive_get_file_metadata(file_id: str) -> Dict[str, Any]:
    """
    Retourne les métadonnées d'un fichier Drive (nom, taille, type, propriétaire, URL).

    Utilise ce tool quand l'utilisateur veut :
    - voir les infos d'un fichier Drive (taille, date, propriétaire)
    - obtenir l'URL directe d'un fichier
    - vérifier les détails d'un document cloud

    Mots-clés : infos fichier, métadonnées, Drive, taille, propriétaire, URL, détails
    """
    svc = get_drive_service()
    try:
        resp = svc.files().get(
            fileId=file_id,
            fields="id,name,mimeType,modifiedTime,size,owners,driveId",
            **_PARTAGES,
        ).execute()
        return {
            "status": "ok",
            "file": {
                "id": resp["id"],
                "name": resp["name"],
                "mimeType": resp.get("mimeType"),
                "size": resp.get("size"),
                "modifiedTime": resp.get("modifiedTime"),
                "owners": [{"email": o.get("emailAddress", "?")}
                           for o in resp.get("owners", [])],
                "drive_partage": bool(resp.get("driveId")),
                "url": _file_url(resp["id"]),
            },
        }
    except HttpError as e:
        if e.resp is not None and e.resp.status == 404:
            return {"status": "not_found", "file_id": file_id,
                    "error": "Fichier introuvable ou accès refusé"}
        return {"status": "error", "error": str(e)}
